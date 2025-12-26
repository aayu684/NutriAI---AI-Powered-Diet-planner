from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_thesis_docx():
    doc = Document()

    # Title Page
    title = doc.add_heading('PROJECT REPORT: AI DIET PLANNER', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n')
    
    p = doc.add_paragraph('Submitted by:')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('[Your Name]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # Acknowledgment
    doc.add_heading('ACKNOWLEDGMENT', level=1)
    doc.add_paragraph(
        "I would like to express my sincere gratitude to my project guide and the department for providing "
        "the opportunity to work on this innovative project. I also thank the creators of the open-source "
        "libraries (Streamlit, Google Generative AI) that made this development possible."
    )
    doc.add_page_break()

    # Abstract
    doc.add_heading('ABSTRACT', level=1)
    doc.add_paragraph(
        "The AI Diet Planner is an intelligent web application designed to address the growing need for "
        "personalized, accessible, and affordable nutrition advice. Traditional diet apps often rely on "
        "static templates that fail to account for individual constraints such as specific allergies, "
        "cultural food preferences (e.g., Indian vegetarian), and complex health goals. This project "
        "leverages Large Language Models (LLM), specifically Google's Gemini 2.5 Flash, to function as "
        "a virtual clinical nutritionist. The system accepts detailed user physiological data, processes "
        "it through a structured reasoning engine, and generates a medically relevant, weekly diet plan. "
        "The output is presented via an interactive dashboard and a downloadable professional PDF. This "
        "report details the system's design, implementation using Python and Agentic AI workflows, and "
        "its potential to revolutionize personal health management."
    )
    doc.add_page_break()

    # Index (Placeholder as Word can generate this, but we'll list sections)
    doc.add_heading('INDEX', level=1)
    index_items = [
        "1. ABOUT THE SYSTEM",
        "   1.1 Problem Definition",
        "   1.2 Requirement Specifications",
        "   1.3 Tools and Technology Used",
        "2. SYSTEM DESIGN USING UML",
        "   2.1 Data Flow Diagrams",
        "   2.2 Use Case Diagram",
        "   2.3 Class Diagram",
        "   2.4 Sequence Diagrams",
        "3. DATA DICTIONARY",
        "   3.1 Data Models",
        "4. IMPLEMENTATION",
        "   4.1 Screen Layouts with Validations",
        "   4.2 Sample Coding",
        "   4.3 Steps of Execution",
        "5. CONCLUSION",
        "   5.1 Importance of Work",
        "   5.2 Future Enhancements",
        "6. REFERENCES"
    ]
    for item in index_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. About the System
    doc.add_heading('1. ABOUT THE SYSTEM', level=1)
    
    doc.add_heading('1.1 PROBLEM DEFINITION (Identification of Needs)', level=2)
    doc.add_paragraph(
        "In the modern world, lifestyle diseases like obesity, diabetes, and hypertension are rising. "
        "While diet is the primary factor in managing these conditions, professional nutritional advice is often:"
    )
    doc.add_paragraph("1. Expensive: Clinical nutritionists charge high consultation fees.", style='List Number')
    doc.add_paragraph("2. Inaccessible: Many people lack access to qualified experts.", style='List Number')
    doc.add_paragraph("3. Generic: Free apps provide 'one-size-fits-all' plans that ignore specific constraints.", style='List Number')
    doc.add_paragraph(
        "There is a need for a system that can 'reason' like a human nutritionist to create truly personalized "
        "plans instantly and at zero marginal cost."
    )

    doc.add_heading('1.2 REQUIREMENT SPECIFICATIONS (Product/System Tasks)', level=2)
    doc.add_paragraph("The system is designed to perform the following tasks:")
    doc.add_paragraph("User Profiling: Capture detailed data including Age, Gender, Height, Weight, Activity Level, Goals, Dietary Restrictions, Allergies, and Cooking Skill.", style='List Bullet')
    doc.add_paragraph("Intelligent Planning: Generate a 7-day meal plan (Breakfast, Lunch, Dinner, Snacks) that strictly adheres to the user's calorie and macro-nutrient needs.", style='List Bullet')
    doc.add_paragraph("Cultural Context: Prioritize Indian-friendly meals when requested or appropriate.", style='List Bullet')
    doc.add_paragraph("Validation: Ensure the output is structured and complete (no missing calories or ingredients).", style='List Bullet')
    doc.add_paragraph("Export: Generate a downloadable PDF report containing the profile summary, weekly schedule, and shopping list.", style='List Bullet')

    doc.add_heading('1.3 TOOLS AND TECHNOLOGY USED', level=2)
    doc.add_paragraph("The project is implemented using a robust Python-based stack:")
    p = doc.add_paragraph()
    p.add_run('Frontend: ').bold = True
    p.add_run('Streamlit (v1.28+) - Used for building the interactive web interface, sidebar forms, and dashboard.')
    
    p = doc.add_paragraph()
    p.add_run('Language: ').bold = True
    p.add_run('Python (v3.13) - Core logic and backend processing.')

    p = doc.add_paragraph()
    p.add_run('AI Service / API: ').bold = True
    p.add_run('Google Generative AI (Gemini 2.5 Flash) - The core intelligence engine used for reasoning and content generation.')

    p = doc.add_paragraph()
    p.add_run('Data Validation: ').bold = True
    p.add_run('Pydantic - Used to define strict data schemas (UserProfile, WeeklyDietPlan) to prevent AI hallucinations and ensure type safety.')

    p = doc.add_paragraph()
    p.add_run('Document Generation: ').bold = True
    p.add_run('ReportLab - Used for programmatic PDF creation (drawing tables, styles, and layouts).')

    p = doc.add_paragraph()
    p.add_run('Environment Management: ').bold = True
    p.add_run('Python-Dotenv - For secure API key management.')

    # 2. System Design
    doc.add_heading('2. SYSTEM DESIGN USING UML', level=1)

    doc.add_heading('2.1 DATA FLOW DIAGRAMS (DFD)', level=2)
    doc.add_paragraph("Level 0 DFD (Context Diagram):")
    doc.add_paragraph("[User] --(User Profile Data)--> [AI Diet Planner System] --(Diet Plan & PDF)--> [User]")
    
    doc.add_paragraph("Level 1 DFD:")
    doc.add_paragraph("1. [User] --Input--> [Profile Capture Module]", style='List Number')
    doc.add_paragraph("2. [Profile Capture Module] --Structured Data--> [AI Processing Engine]", style='List Number')
    doc.add_paragraph("3. [AI Processing Engine] --Prompt + Schema--> [Gemini API]", style='List Number')
    doc.add_paragraph("4. [Gemini API] --JSON Response--> [Validation Engine]", style='List Number')
    doc.add_paragraph("5. [Validation Engine] --Validated Object--> [Dashboard UI]", style='List Number')
    doc.add_paragraph("6. [Validation Engine] --Object--> [PDF Generator] --PDF File--> [User]", style='List Number')

    doc.add_heading('2.2 USE CASE DIAGRAM', level=2)
    doc.add_paragraph("Actors: User, AI System")
    doc.add_paragraph("Use Cases:")
    doc.add_paragraph("1. Manage Profile: User enters/updates age, weight, goals.", style='List Bullet')
    doc.add_paragraph("2. Generate Plan: User triggers the AI generation process.", style='List Bullet')
    doc.add_paragraph("3. View Dashboard: User views daily meals and nutritional metrics.", style='List Bullet')
    doc.add_paragraph("4. Download Report: User downloads the plan as a PDF.", style='List Bullet')

    doc.add_heading('2.3 CLASS DIAGRAM', level=2)
    doc.add_paragraph("The system is built around the following key classes:")
    doc.add_paragraph("UserProfile: Attributes include name, age, weight_kg, goal, dietary_restrictions.", style='List Bullet')
    doc.add_paragraph("WeeklyDietPlan: Attributes include user_profile, daily_plans, weekly_summary.", style='List Bullet')
    doc.add_paragraph("DailyPlan: Attributes include day, meals, total_calories.", style='List Bullet')
    doc.add_paragraph("MealPlan: Attributes include meal_name, ingredients, nutrition_info.", style='List Bullet')
    doc.add_paragraph("AIDietitian: Methods include create_diet_plan(UserProfile).", style='List Bullet')
    doc.add_paragraph("DietPlanPDFGenerator: Methods include generate_diet_plan_pdf(WeeklyDietPlan).", style='List Bullet')

    doc.add_heading('2.4 SEQUENCE DIAGRAMS', level=2)
    doc.add_paragraph("Scenario: Generating a Diet Plan")
    steps = [
        "User enters details in Sidebar UI.",
        "UI creates UserProfile instance.",
        "User clicks 'Generate'.",
        "UI calls AIDietitian.create_diet_plan(profile).",
        "AIDietitian constructs prompt and calls Gemini API.",
        "Gemini API returns JSON.",
        "AIDietitian validates JSON into WeeklyDietPlan object.",
        "AIDietitian returns object to UI.",
        "UI renders Dashboard.",
        "User clicks 'Download PDF'.",
        "UI calls DietPlanPDFGenerator.generate().",
        "PDFGenerator returns file path.",
        "UI serves file to User."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # 3. Data Dictionary
    doc.add_heading('3. DATA DICTIONARY', level=1)
    doc.add_heading('3.1 DATA MODELS', level=2)
    doc.add_paragraph("Since the application is stateless and uses AI generation on-the-fly, a traditional SQL database is not used. Instead, strict Pydantic Data Models act as the schema.")
    
    doc.add_paragraph("Table 1: UserProfile Schema").bold = True
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    data = [
        ('name', 'String', 'Name of the user'),
        ('age', 'Integer', 'Age in years'),
        ('gender', 'String', 'Male/Female/Other'),
        ('weight_kg', 'Float', 'Current weight'),
        ('activity_level', 'Enum', 'Sedentary, Lightly Active, etc.'),
        ('goal', 'Enum', 'Weight Loss, Muscle Gain, etc.'),
        ('allergies', 'List[String]', 'E.g., ["Peanuts", "Shellfish"]')
    ]
    for field, dtype, desc in data:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = dtype
        row_cells[2].text = desc

    doc.add_paragraph("\n")
    doc.add_paragraph("Table 2: MealPlan Schema").bold = True
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Field'
    hdr_cells2[1].text = 'Type'
    hdr_cells2[2].text = 'Description'
    
    data2 = [
        ('meal_time', 'Enum', 'Breakfast, Lunch, Dinner, Snack'),
        ('meal_name', 'String', 'Name of the dish'),
        ('calories', 'Integer', 'Energy content'),
        ('protein', 'Float', 'Protein in grams'),
        ('ingredients', 'List[String]', 'List of required items')
    ]
    for field, dtype, desc in data2:
        row_cells = table2.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = dtype
        row_cells[2].text = desc

    # 4. Implementation
    doc.add_heading('4. IMPLEMENTATION', level=1)
    
    doc.add_heading('4.1 SCREEN LAYOUTS WITH VALIDATIONS', level=2)
    doc.add_paragraph("The application uses a Sidebar Layout for inputs to keep the main area clean for results.")
    doc.add_paragraph("Validations:", style='List Bullet')
    doc.add_paragraph("Numeric Inputs: Age (1-120), Weight (20-500kg) are constrained by Streamlit widgets.", style='List Bullet')
    doc.add_paragraph("Selection Inputs: Goals and Activity Levels are restricted to Dropdowns (Enums) to prevent invalid text entry.", style='List Bullet')
    doc.add_paragraph("Schema Validation: The Pydantic library validates the AI's output.", style='List Bullet')

    doc.add_heading('4.2 SAMPLE CODING', level=2)
    doc.add_paragraph("Core Logic (ai_dietitian.py):")
    code_sample = """
def create_diet_plan(self, user_profile: UserProfile) -> Optional[WeeklyDietPlan]:
    # Construct the prompt with Chain-of-Thought instructions
    prompt = (
        f"{self.cot_prompts['meal_planning']}\\n\\n"
        "User profile:\\n"
        f"{user_profile.model_dump_json(indent=2)}\\n\\n"
        "Create a realistic, budget-aware weekly diet plan..."
    )

    # Call Gemini with Structured Output enforcement
    model = genai.GenerativeModel(self.model_name)
    resp = model.generate_content(
        prompt,
        generation_config=genai.GenerationConfig(
            response_mime_type="application/json",
            response_schema=WeeklyDietPlan, # Passing the class directly
        ),
    )
    # Parse into object
    return WeeklyDietPlan(**json.loads(resp.text))
    """
    doc.add_paragraph(code_sample, style='Quote')

    doc.add_heading('4.3 STEPS OF EXECUTION', level=2)
    steps = [
        "Setup: Install dependencies via `pip install -r requirements.txt`.",
        "Configuration: Create `.env` file with `GEMINI_API_KEY`.",
        "Launch: Run command `streamlit run app.py`.",
        "Navigation:",
        "   - Step 1: Fill 'User Profile' in the left sidebar.",
        "   - Step 2: Select 'Goals & Lifestyle'.",
        "   - Step 3: Click 'Generate My Diet Plan' in the main area.",
        "   - Step 4: View the 'Weekly Overview' metrics.",
        "   - Step 5: Expand 'Daily Meal Plans' to see details.",
        "   - Step 6: Click 'Generate PDF' to download the report."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # 5. Conclusion
    doc.add_heading('5. CONCLUSION', level=1)
    
    doc.add_heading('5.1 IMPORTANCE OF THE WORKED-OUT TASKS', level=2)
    doc.add_paragraph(
        "This project successfully demonstrates that Agentic AI can replace static rule-based systems in complex "
        "domains like healthcare and nutrition. By combining Structured Output generation with Domain-Specific "
        "Prompting, we achieved a system that produces high-quality, error-free, and personalized results. "
        "The use of Pydantic ensures that the AI is not just a chatbot, but a reliable data generator capable "
        "of integrating with software systems (PDF generators, databases)."
    )

    doc.add_heading('5.2 FURTHER ENHANCEMENT IN FUTURE', level=2)
    doc.add_paragraph("1. Image Integration: Using models like Imagen 3 to generate realistic photos of the suggested meals.", style='List Number')
    doc.add_paragraph("2. Grocery APIs: Connecting the shopping list directly to delivery apps.", style='List Number')
    doc.add_paragraph("3. User Accounts: Adding a database to track user progress.", style='List Number')
    doc.add_paragraph("4. Feedback Loop: Allowing users to 'reject' a specific meal and have the AI regenerate just that slot.", style='List Number')

    # 6. References
    doc.add_heading('6. REFERENCES', level=1)
    refs = [
        "Google AI Studio Documentation: For Gemini API and Structured Output.",
        "Streamlit Documentation: For UI components and Session State.",
        "Pydantic Documentation: For data modeling and validation.",
        "ReportLab User Guide: For PDF generation techniques.",
        "Research: 'Prompt Engineering for Large Language Models' - Techniques on Chain-of-Thought prompting."
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='List Bullet')

    doc.save('Thesis_Report.docx')
    print("Thesis_Report.docx created successfully.")

if __name__ == "__main__":
    create_thesis_docx()
